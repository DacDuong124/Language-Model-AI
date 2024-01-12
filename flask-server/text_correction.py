import requests
from docx import Document
from urllib.parse import quote, urlencode, urlparse, unquote
import time
import os
import shutil
from firebase_admin import credentials, auth, firestore, storage
from google.cloud import storage
BATCH_SIZE = 3
# Configure Firebase Admin SDK
cred = credentials.Certificate("language-ai-model-firebase-adminsdk-l4hgq-1c59e87bd8.json")


def custom_urlencode(params):
    return urlencode(params, quote_via=quote)

def left_whitespaces_counter(string):
    return len(string) - len(str(string).lstrip())


def download_document(url):
    try:
        # Decode the URL
        decoded_url = unquote(url)

        # Extracting the filename from the decoded URL
        parsed_url = urlparse(decoded_url)
        filename = os.path.basename(parsed_url.path)

        # Check if the URL contains a valid filename
        if not filename:
            raise ValueError("URL does not contain a valid filename")

        response = requests.get(url)
        response.raise_for_status()

        with open(filename, 'wb') as f:
            f.write(response.content)
        return filename  # Return the actual filename used
    except Exception as e:
        print(f"Failed to download document: {e}")
        return None

def correct_document_from_url(document_url, user_id):
    # Download the document and get the filename
    downloaded_filename = download_document(document_url)
    if not downloaded_filename:
        return None  # Exit if the download fails

    local_corrected_path = None
    file_extension = os.path.splitext(downloaded_filename)[1].lower()

    if file_extension == '.docx':
        doc = Document(downloaded_filename)
        paragraphs_batch = []
        replace_list = {}

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs_batch.append(text)

                if len(paragraphs_batch) >= BATCH_SIZE:
                    batch_replacements = call_batch(paragraphs_batch)
                    if batch_replacements:
                        replace_list.update(batch_replacements)
                    paragraphs_batch = []

        if paragraphs_batch:  # Process any remaining paragraphs
            batch_replacements = call_batch(paragraphs_batch)
            if batch_replacements:
                replace_list.update(batch_replacements)

        if replace_list:
            change_sentences(replace_list, downloaded_filename.rstrip(file_extension))
            change_format(downloaded_filename.rstrip(file_extension), downloaded_filename.rstrip(file_extension) + "_raw")

            local_corrected_path = os.path.splitext(downloaded_filename)[0] + "_corrected.docx"
            shutil.move(downloaded_filename.rstrip(file_extension) + "_raw.docx", local_corrected_path)

    elif file_extension == '.txt':
        lines = read_txt_file(downloaded_filename)
        lines_batch = []
        corrected_lines = []

        for line in lines:
            stripped_line = line.strip()
            if stripped_line:
                lines_batch.append(stripped_line)

                if len(lines_batch) >= BATCH_SIZE:
                    batch_replacements = call_batch(lines_batch)
                    for original, corrected in zip(lines_batch, batch_replacements.values()):
                        corrected_lines.append(corrected + '\n' if corrected else original + '\n')
                    lines_batch = []

        if lines_batch:  # Process any remaining lines
            batch_replacements = call_batch(lines_batch)
            for original, corrected in zip(lines_batch, batch_replacements.values()):
                corrected_lines.append(corrected + '\n' if corrected else original + '\n')

        local_corrected_path = os.path.splitext(downloaded_filename)[0] + "_corrected.txt"
        write_txt_file(local_corrected_path, corrected_lines)

    if local_corrected_path:
        # Upload the corrected file
        corrected_file_url = upload_to_firebase(local_corrected_path, user_id)
        if corrected_file_url:
            corrected_file_name = os.path.basename(local_corrected_path)
            add_file_url_to_firestore(corrected_file_url, corrected_file_name, user_id)
        else:
            print("Failed to upload corrected file to Firebase Storage")

        # Check if there is a formatted file for .docx, and upload it
        if file_extension == '.docx':
            formatted_file_path = os.path.splitext(downloaded_filename)[0] + "_formatted.docx"
            formatted_file_url = upload_to_firebase(formatted_file_path, user_id)
            if formatted_file_url:
                formatted_file_name = os.path.basename(formatted_file_path)
                add_file_url_to_firestore(formatted_file_url, formatted_file_name, user_id)
            else:
                print("Failed to upload formatted file to Firebase Storage")

        # Delete the original downloaded file and any intermediate files
        os.remove(downloaded_filename)
        if file_extension == '.docx':
            raw_file_path = downloaded_filename.rstrip(file_extension) + "_raw.docx"
            if os.path.exists(raw_file_path):
                os.remove(raw_file_path)
            if os.path.exists(formatted_file_path):
                os.remove(formatted_file_path)

        # Delete the corrected file
        if os.path.exists(local_corrected_path):
            os.remove(local_corrected_path)

        return corrected_file_url, formatted_file_url  # Return both URLs
    else:
        print("No corrected file to upload.")
        os.remove(downloaded_filename)
        return None
    
def process_document(doc_name):
    doc = Document(doc_name + ".docx")
    paragraphs_batch = []
    replace_list = {}

    # Iterate through paragraphs and batch them
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs_batch.append(text)

            # If batch size is reached, process the batch
            if len(paragraphs_batch) >= BATCH_SIZE:
                batch_replacements = call_batch(paragraphs_batch)
                if batch_replacements:
                    replace_list.update(batch_replacements)
                paragraphs_batch = []  # Reset the batch for next set of paragraphs

    # Process any remaining paragraphs in the final batch
    if paragraphs_batch:
        batch_replacements = call_batch(paragraphs_batch)
        if batch_replacements:
            replace_list.update(batch_replacements)

    # Apply replacements to the document regardless of whether there are replacements
    change_sentences(replace_list, doc_name.rstrip('.docx'))
    change_format(doc_name.rstrip('.docx'), doc_name.rstrip('.docx') + "_raw")

    # Move the corrected file to a new path
    local_corrected_path = doc_name.rstrip('.docx') + "_corrected.docx"
    shutil.move(doc_name.rstrip('.docx') + "_raw.docx", local_corrected_path)
    return local_corrected_path



# Initialize a storage client
storage_client = storage.Client()
# Get the bucket from the storage client
bucket = storage_client.get_bucket('language-ai-model.appspot.com')
def upload_to_firebase(file_path, user_id):
    try:
        # Create a blob in the specific user's folder
        blob = bucket.blob(f'user_files/{user_id}/{os.path.basename(file_path)}')

        # Upload the file
        blob.upload_from_filename(file_path)

        # Make the blob publicly accessible (if your policy allows this)
        blob.make_public()

        # Return the public URL for the file
        return blob.public_url
    except Exception as e:
        print(f"Failed to upload file to Firebase: {e}")
        return None


def add_file_url_to_firestore(file_url, file_name, user_id):
    try:
        # Initialize Firestore DB
        db = firestore.client()

        # Create a new document in the user's 'documents' subcollection
        file_data = {
            'user_id': user_id,
            'url': file_url,  # URL of the corrected file
            'name': file_name, # Include the file name
            'status': 'active',  
            'createdAt': firestore.SERVER_TIMESTAMP  # Use a timestamp to track when the file was added
        }
        # Add the document to the 'documents' subcollection of the specific user
        # db.collection('users').document(user_id).collection('documents').add(file_data)
        doc_ref = db.collection('users').document(user_id).collection('documents').document(file_name)
        doc_ref.set(file_data)
        print("Corrected file URL added to Firestore in the user's 'documents' subcollection successfully.")
    except Exception as e:
        print(f"Failed to add corrected file URL to Firestore: {e}")







def call_batch(paragraphs):
    prompt_questions = ["Correct this text: "]
    additional_prompt = " Here are the corrected versions:"

    # Combine the paragraphs with a delimiter for processing
    combined_input = "\n\n".join(paragraphs)
    input_prompt = combined_input + '.' if not combined_input.endswith('.') else combined_input

    replacements = {}

    for question in prompt_questions:
        params = {'max_length': len(combined_input), 'prompts': question + input_prompt + additional_prompt}
        response = requests.get("https://polite-horribly-cub.ngrok-free.app/generate_code", params=custom_urlencode(params))
        print("Request URL:", response.url)  # Debugging

        if response.status_code == 200:
            try:
                generated_code_list = response.json()
            except Exception as e:
                print("Failed to parse JSON response:", e)
                return None

            for res in generated_code_list:
                formatted = (str(res).split("\n\n")[0].replace("\n         ", '').split("\n        ")[0].replace('\r', '').strip())

                # Split the corrected paragraphs and map them back to the original paragraphs
                corrected_paragraphs = formatted.split("\n\n")
                for original, corrected in zip(paragraphs, corrected_paragraphs):
                    replacements[original] = corrected

                # More debugging information
                print("Formatted Response:", formatted)
                return replacements

        else:
            print("Failed to retrieve code. Status code:", response.status_code)
            return None

def change_sentences(replacements, original):
    if not replacements:
        print("No replacements found")
        return

    og = Document(original + ".docx")

    for og_para in og.paragraphs:
        for target_sentence, new_text in replacements.items():
            if new_text is None or target_sentence not in og_para.text:
                continue
            print(f"Replacing '{target_sentence}' with '{new_text}'")  # Debugging
            og_para.text = og_para.text.replace(target_sentence, new_text)

    og.save(original + "_raw.docx")


def change_format(source, output):
    source_doc = Document(source + ".docx")
    output_doc = Document(output + ".docx")

    for source_paragraph, output_paragraph in zip(source_doc.paragraphs, output_doc.paragraphs):
        output_paragraph.paragraph_format.alignment = source_paragraph.paragraph_format.alignment
        output_paragraph.paragraph_format.first_line_indent = source_paragraph.paragraph_format.first_line_indent
        output_paragraph.paragraph_format.left_indent = source_paragraph.paragraph_format.left_indent
        output_paragraph.paragraph_format.space_before = source_paragraph.paragraph_format.space_before
        output_paragraph.paragraph_format.space_after = source_paragraph.paragraph_format.space_after

        for source_run, output_run in zip(source_paragraph.runs, output_paragraph.runs):
            output_run.bold = source_run.bold
            output_run.italic = source_run.italic
            output_run.underline = source_run.underline
            output_run.font.name = source_run.font.name
            output_run.font.size = source_run.font.size
            output_run.font.color.rgb = source_run.font.color.rgb

    output_doc.save(source + "_formatted.docx")

def read_txt_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.readlines()

def write_txt_file(file_path, lines):
    with open(file_path, 'w', encoding='utf-8') as file:
        file.writelines(lines)

def correct_text_file(file_path):
    extension = os.path.splitext(file_path)[1]
    if extension.lower() == '.docx':
        return correct_document_from_url(file_path)
    elif extension.lower() == '.txt':
        lines = read_txt_file(file_path)
        corrected_lines = []
        for line in lines:
            if line.strip():
                corrected_line = call_batch(line.strip())
                corrected_lines.append(corrected_line + '\n' if corrected_line else line)
            else:
                corrected_lines.append('\n')
        corrected_file_path = file_path.replace('.txt', '_corrected.txt')
        write_txt_file(corrected_file_path, corrected_lines)
        return corrected_file_path
    else:
        print("Unsupported file format")
        return None


if __name__ == "__main__":
    start = time.perf_counter()
    doc_name = "..."
    doc = Document(doc_name + ".docx")
    replace_list = {}
    
    for paragraph in doc.paragraphs:
        if str(paragraph.text).strip() != '':
            target = paragraph.text.strip()
            answer = call_batch(target)
            if answer:
                replace_list[target] = answer
                print(f'{target} ---> {answer}')
            else:
                print(f"No replacement for '{target}'")

    if replace_list:
        change_sentences(replace_list, doc_name)
        change_format(doc_name, doc_name + "_raw")
    else:
        print("No replacements made.")

    end = time.perf_counter()
    print("\n")
    print(f'Corrections completed in {end - start} seconds')